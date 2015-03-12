from firebase import firebase
from docx import Document
from datetime import date
import re

#**************AUTHENTICATION IS NEEDED TO MAKE SURE ONLY WE CAN READ THIS
firebase = firebase.FirebaseApplication('https://electionhack.firebaseio.com', None)
result = firebase.get('/forms/testform', None)

stuff = {
    'addr_city': '',
    'addr_first': '',
    'addr_second': '',
    'firstname': '',
    'DoB_d': '',
    'DoB_m': '',
    'DoB_y': '',
    'commonforename': '',
    'commonsurname': '',
    'othernames': '',
    'surname': '',
    'addr_postcode': '',
    'email': '',
    'election_date': '07/05/2015',
    'constituency': '',
    'today': date.today().strftime('%d/%m/%Y')
}

stuff.update(result);

##turn keys into placeholders
templateStuff = dict(('XXX%sXXX' % k, str(v)) for (k, v) in stuff.iteritems())


infile = 'nominationpapers.docx'
outfile = 'testform-complete.docx'

def multiple_find(vals, text):
    """ find any of the dictionary keys in the text""" 

    # Create a regular expression  from the dictionary keys
    regex = re.compile("(%s)" % "|".join(map(re.escape, vals.keys())))

    # For each match, look-up corresponding value in dictionary
    return regex.search(text) 

def multiple_replace(vals, text): 

    """ Replace in 'text' all occurences of any key in the given
    dictionary by its corresponding value.  Returns the new tring.""" 

    # Create a regular expression  from the dictionary keys
    regex = re.compile("(%s)" % "|".join(map(re.escape, vals.keys())))

    # For each match, look-up corresponding value in dictionary
    return regex.sub(lambda mo: vals[mo.string[mo.start():mo.end()]], text) 

form = Document('nominationpapers.docx')
for table in form.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text
                if multiple_find(templateStuff, text):
                    paragraph.clear()
                    paragraph.add_run(text=multiple_replace(templateStuff, text))

form.save('testform.docx')
